from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('PhonePe Transaction Insights – Project Report', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('')

# --- SECTION 1: How I approached this ---
doc.add_heading('How I Approached This', level=1)

p = doc.add_paragraph(
    "I picked PhonePe transaction data because honestly, payments data from India is something I find genuinely interesting. "
    "The way people transact across different states and categories tells a lot more than just numbers — it reflects spending habits, "
    "regional economic patterns, and how digital payments have grown in this country over the last few years."
)

p = doc.add_paragraph(
    "I started by just pulling the raw data and looking at it. The dataset had transaction records across 36 Indian states and union territories, "
    "and the first thing I noticed was that the data wasn't clean — there were inconsistencies in state names, "
    "some null values in the category columns, and the date fields needed parsing before I could do anything meaningful with them. "
    "So before any analysis, I spent a good chunk of time on the ETL side."
)

p = doc.add_paragraph(
    "Once the data was clean, I loaded it into a PostgreSQL database. I designed 9 tables keeping normalization in mind — "
    "things like separating transactions, users, devices, and location data so queries wouldn't be slow or messy. "
    "The total data I was working with was around 150MB which isn't massive but it was enough to make poor schema design "
    "noticeably painful, so I was careful about indexing and joins from the start."
)

doc.add_paragraph('')

# --- SECTION 2: What I actually did ---
doc.add_heading('What I Actually Did', level=1)

p = doc.add_paragraph(
    "Most of the work fell into three areas — SQL analysis, exploratory data analysis in Python, and building the dashboard."
)

p = doc.add_paragraph(
    "For the SQL part, I wrote 9 business-focused queries. These weren't just SELECT * type things — "
    "I was looking at things like which states had the highest transaction volumes, what the average transaction value looked like "
    "across different payment categories, and how growth trended quarter over quarter. "
    "Writing these queries made me think like someone in a business intelligence role rather than just a developer."
)

p = doc.add_paragraph(
    "The EDA part in Python was where I spent the most time. I built 22 visualizations in total — "
    "covering time-series trends, geo maps showing state-level distribution, category breakdowns, and user-device analysis. "
    "I used matplotlib and seaborn mostly, and for the geo plots I used folium. "
    "Some charts didn't turn out useful the first time so I redid them — the geo heatmap especially took a few tries to make readable."
)

p = doc.add_paragraph(
    "The dashboard was built in Streamlit and deployed to Streamlit Cloud. "
    "I wanted it to actually be interactive, not just static images, so I added filters for state, date range, and category. "
    "The goal was that someone with no coding background should be able to open it and draw conclusions within a few minutes."
)

doc.add_paragraph('')

# --- SECTION 3: Insights I Found ---
doc.add_heading('Insights I Found', level=1)

p = doc.add_paragraph(
    "A few things stood out to me during the analysis that I wasn't expecting going in."
)

p = doc.add_paragraph(
    "First — Maharashtra, Karnataka, and Telangana together accounted for a disproportionately large share of total transaction volume. "
    "This wasn't surprising on its own but what was interesting is that when I looked at average transaction value, "
    "smaller states like Goa and Sikkim had higher per-transaction amounts. "
    "So volume and value don't always move together — which has real implications for how you'd prioritize merchant acquisition if you were a fintech."
)

p = doc.add_paragraph(
    "Second — merchant payments and recharges/bill payments dominated the category breakdown. "
    "Peer-to-peer transfers were high in frequency but lower in average value, which makes sense. "
    "What I found interesting was the spike in bill payment transactions in the March-April window every year — "
    "likely driven by end-of-financial-year utility settlements and insurance renewals."
)

p = doc.add_paragraph(
    "Third — Android devices completely dominated across all states with very little variation. "
    "But in metro cities the iOS share was noticeably higher than the national average. "
    "This kind of device segmentation is something a product team would actually care about for UI prioritization."
)

p = doc.add_paragraph(
    "Fourth — transaction growth was not linear. There were specific quarters that showed sharp jumps, "
    "likely tied to cashback campaigns or government initiatives like DPDP or UPI incentive programs. "
    "The time-series analysis made these inflection points very visible."
)

doc.add_paragraph('')

# --- SECTION 4: What Could Be Done Next ---
doc.add_heading('What Could Be Done Next', level=1)

p = doc.add_paragraph(
    "There's a lot of room to take this further and I have a few directions in mind."
)

p = doc.add_paragraph(
    "The most obvious next step is adding predictive modeling — using the historical transaction trends to forecast "
    "state-level or category-level volumes for the next quarter. A simple ARIMA or Prophet model would work here and "
    "wouldn't be overkill for the data size."
)

p = doc.add_paragraph(
    "Another thing I'd like to add is anomaly detection. Right now the dashboard just shows you what happened, "
    "but it'd be much more useful if it could flag unusual spikes or drops automatically. "
    "Something like an Isolation Forest or even a Z-score based alert on rolling averages could work well."
)

p = doc.add_paragraph(
    "I also want to bring in external data — things like state-level GDP, internet penetration rates, and population density — "
    "and see how well PhonePe adoption correlates with those. That would make the geo analysis a lot richer "
    "and move it from descriptive to actually explanatory."
)

p = doc.add_paragraph(
    "On the dashboard side, I'd like to add a user cohort analysis feature — tracking how transaction behavior "
    "changes for the same user segment over time. That's something that's hard to do with aggregate data "
    "but would be possible with a more granular dataset."
)

doc.add_paragraph('')

# --- SECTION 5: Personal Takeaway ---
doc.add_heading('What I Took Away From This', level=1)

p = doc.add_paragraph(
    "Honestly this project taught me more about how to think about data than about any specific tool. "
    "The SQL schema design forced me to think ahead about how the data would be queried. "
    "The EDA made me realize how easy it is to build charts that look nice but don't say anything — "
    "I had to keep asking myself 'so what does this actually tell someone?'"
)

p = doc.add_paragraph(
    "The dashboard deployment was also a good reminder that analysis only has value if someone can actually use it. "
    "A notebook on your local machine helps nobody. Getting it live on Streamlit Cloud, with filters and real data, "
    "is what makes it a real project rather than just an exercise."
)

p = doc.add_paragraph(
    "If I were to do it again I'd probably start the schema design on paper before touching any code. "
    "I made a couple of structural decisions early that I had to backtrack on, and that cost time. "
    "But overall I'm happy with where it landed."
)

# Save
doc.save("PhonePe_Project_Report.docx")
print("Done")